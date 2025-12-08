#!/usr/bin/env python3
"""
Synthetic HTTP MCP Server

A single HTTP server that masquerades as multiple MCP servers.
Routes requests based on path: /mcp/{server_name}/tools/{tool_name}

Handler types:
- filesystem: Read files from the synthetic data directory
- table_lookup: Look up data by key from JSON files
- excel_reader: Read Excel files
- static_json: Return fixed JSON responses
- compute: Perform simple computations

Stub logging:
When a static_json handler returns a stub response (containing "synthetic stub"),
the call is logged to synthetic_mcp/logs/stub_calls.json for later review.
"""

import csv
import json
import threading
from datetime import datetime
from http.server import BaseHTTPRequestHandler, HTTPServer
from pathlib import Path
from typing import Any
from urllib.parse import urlparse

# Default paths - can be overridden
DEFAULT_CONFIG_PATH = Path(__file__).parent.parent / "config" / "servers.json"
DEFAULT_DATA_PATH = Path(__file__).parent.parent / "data"
DEFAULT_LOGS_PATH = Path(__file__).parent.parent / "logs"


class StubCallLogger:
    """Thread-safe logger for stub handler calls.
    
    Logs each stub call with server name, tool name, and parameters.
    Aggregates calls so each unique (server, tool) pair appears once
    with a count and example parameters.
    """
    
    def __init__(self, log_path: Path):
        self.log_path = log_path
        self.lock = threading.Lock()
        self.calls: dict[tuple[str, str], dict] = {}
        self._load_existing()
    
    def _load_existing(self) -> None:
        """Load existing log file if it exists."""
        if self.log_path.exists():
            try:
                with open(self.log_path, encoding="utf-8") as f:
                    data = json.load(f)
                    for entry in data.get("stub_calls", []):
                        key = (entry["server"], entry["tool"])
                        self.calls[key] = entry
            except (json.JSONDecodeError, KeyError):
                pass
    
    def log_call(
        self,
        server_name: str,
        tool_name: str,
        params: dict,
        response: Any,
    ) -> None:
        """Log a stub handler call."""
        with self.lock:
            key = (server_name, tool_name)
            
            if key not in self.calls:
                self.calls[key] = {
                    "server": server_name,
                    "tool": tool_name,
                    "call_count": 0,
                    "first_seen": datetime.now().isoformat(),
                    "example_params": [],
                    "stub_response": response,
                    "annotation": "",  # For manual review
                    "handler_override": None,  # For specifying a real handler
                }
            
            entry = self.calls[key]
            entry["call_count"] += 1
            entry["last_seen"] = datetime.now().isoformat()
            
            # Keep up to 3 example parameter sets
            if len(entry["example_params"]) < 3 and params not in entry["example_params"]:
                entry["example_params"].append(params)
            
            self._save()
    
    def _save(self) -> None:
        """Save the log to disk."""
        self.log_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Sort by call count (most called first)
        sorted_calls = sorted(
            self.calls.values(),
            key=lambda x: x["call_count"],
            reverse=True
        )
        
        data = {
            "_comment": "Stub handler calls - review and add annotations/handler_override as needed",
            "_last_updated": datetime.now().isoformat(),
            "stub_calls": sorted_calls,
        }
        
        with open(self.log_path, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
    
    def clear(self) -> None:
        """Clear all logged calls."""
        with self.lock:
            self.calls = {}
            if self.log_path.exists():
                self.log_path.unlink()


class SyntheticMCPHandler(BaseHTTPRequestHandler):
    """HTTP request handler for synthetic MCP server."""

    servers_config: dict = {}
    data_path: Path = DEFAULT_DATA_PATH
    web_corpus_metadata: dict = {}  # Cached web corpus metadata
    decoy_urls: list = []  # Cached decoy URLs for search results
    stub_logger: StubCallLogger | None = None  # Optional stub call logger
    current_server_name: str = ""  # Track current server for logging

    def log_message(self, format: str, *args: Any) -> None:
        """Override to reduce logging noise."""
        pass  # Suppress default logging

    def _is_stub_response(self, response: Any) -> bool:
        """Check if a response is from a stub handler."""
        if isinstance(response, dict):
            # Check for stub marker in message field
            message = response.get("message", "")
            if isinstance(message, str) and "synthetic stub" in message.lower():
                return True
        return False

    def send_json_response(self, data: Any, status: int = 200) -> None:
        """Send a JSON response."""
        self.send_response(status)
        self.send_header("Content-Type", "application/json")
        self.end_headers()
        self.wfile.write(json.dumps(data).encode("utf-8"))

    def send_error_response(self, message: str, status: int = 400) -> None:
        """Send an error response."""
        self.send_json_response({"error": message}, status)

    def do_GET(self) -> None:
        """Handle GET requests - list servers or tools."""
        parsed = urlparse(self.path)
        path_parts = parsed.path.strip("/").split("/")

        if path_parts == [""]:
            # Root - list all servers
            servers = list(self.servers_config.keys())
            self.send_json_response({"servers": servers})

        elif len(path_parts) == 2 and path_parts[0] == "mcp":
            # /mcp/{server_name} - list tools
            server_name = path_parts[1]
            if server_name not in self.servers_config:
                self.send_error_response(f"Unknown server: {server_name}", 404)
                return

            server = self.servers_config[server_name]
            tools = [
                {"name": t["name"], "description": t.get("description", "")}
                for t in server.get("tools", [])
            ]
            self.send_json_response({"server": server_name, "tools": tools})

        else:
            self.send_error_response("Invalid path", 404)

    def do_POST(self) -> None:
        """Handle POST requests - execute tools."""
        parsed = urlparse(self.path)
        path_parts = parsed.path.strip("/").split("/")

        # Expected: /mcp/{server_name}/tools/{tool_name}
        if len(path_parts) != 4 or path_parts[0] != "mcp" or path_parts[2] != "tools":
            self.send_error_response(
                "Invalid path. Expected: /mcp/{server}/tools/{tool}", 400
            )
            return

        server_name = path_parts[1]
        tool_name = path_parts[3]

        # Find server and tool
        if server_name not in self.servers_config:
            self.send_error_response(f"Unknown server: {server_name}", 404)
            return

        server = self.servers_config[server_name]
        tool = None
        for t in server.get("tools", []):
            if t["name"] == tool_name:
                tool = t
                break

        if not tool:
            self.send_error_response(
                f"Unknown tool: {tool_name} in server {server_name}", 404
            )
            return

        # Parse request body
        content_length = int(self.headers.get("Content-Length", 0))
        body = (
            self.rfile.read(content_length).decode("utf-8")
            if content_length > 0
            else "{}"
        )

        try:
            params = json.loads(body) if body else {}
        except json.JSONDecodeError:
            self.send_error_response("Invalid JSON in request body", 400)
            return

        # Execute handler
        handler = tool.get("handler", {})
        handler_type = handler.get("type", "static_json")

        try:
            result = self.execute_handler(
                handler_type, handler, params, tool, server_name
            )
            self.send_json_response({"result": result})
        except Exception as e:
            self.send_error_response(f"Handler error: {e!s}", 500)

    def execute_handler(
        self,
        handler_type: str,
        handler: dict,
        params: dict,
        tool: dict,
        server_name: str = "",
    ) -> Any:
        """Execute a handler based on its type."""
        if handler_type == "static_json":
            response = handler.get("response", {})
            
            # Check if this is a stub response and log it
            if self.stub_logger and self._is_stub_response(response):
                self.stub_logger.log_call(
                    server_name=server_name,
                    tool_name=tool.get("name", ""),
                    params=params,
                    response=response,
                )
            
            return response

        elif handler_type == "compute":
            return self.handle_compute(handler, params)

        elif handler_type == "filesystem":
            return self.handle_filesystem(handler, params, tool)

        elif handler_type == "table_lookup":
            return self.handle_table_lookup(handler, params)

        elif handler_type == "excel_reader":
            return self.handle_excel_reader(handler, params, tool)

        elif handler_type == "web_corpus":
            return self.handle_web_corpus(handler, params, tool)

        elif handler_type == "url_search":
            return self.handle_url_search(handler, params, tool)

        elif handler_type == "hackernews_story":
            return self.handle_hackernews_story(handler, params, tool)

        elif handler_type == "wikipedia_search":
            return self.handle_wikipedia_search(handler, params, tool)

        elif handler_type == "table_search":
            return self.handle_table_search(handler, params, tool)

        else:
            raise ValueError(f"Unknown handler type: {handler_type}")

    def handle_compute(self, handler: dict, params: dict) -> Any:
        """Handle compute operations."""
        operation = handler.get("operation", "")

        if operation == "fixed_value":
            return handler.get("value")

        raise ValueError(f"Unknown compute operation: {operation}")

    def handle_filesystem(self, handler: dict, params: dict, tool: dict) -> Any:
        """Handle filesystem operations (read_file, list_directory, etc.)."""
        tool_name = tool.get("name", "")

        # Map /root paths to our synthetic data directory
        files_root = self.data_path / "files" / "root"

        if tool_name in ("read_file", "document_reader", "get_document_text"):
            # Get the path parameter (different tools use different param names)
            file_path = (
                params.get("path")
                or params.get("filePath")
                or params.get("filename", "")
            )

            # Handle head/tail parameters for read_file
            head = params.get("head")
            tail = params.get("tail")

            # Normalize path - remove /root prefix and resolve
            if file_path.startswith("/root/"):
                rel_path = file_path[6:]  # Remove /root/
            elif file_path.startswith("/"):
                rel_path = file_path[1:]
            else:
                rel_path = file_path

            actual_path = files_root / rel_path

            if not actual_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")

            # Handle different file types
            suffix = actual_path.suffix.lower()

            if suffix == ".docx":
                # For Word documents, return extracted text
                return self.read_docx(actual_path)

            elif suffix == ".csv":
                # For CSV files, return as raw text (models can parse it)
                with open(actual_path, encoding="utf-8", errors="replace") as f:
                    lines = f.readlines()
                if head is not None:
                    lines = lines[: int(head)]
                elif tail is not None:
                    lines = lines[-int(tail) :]
                return "".join(lines)

            else:
                # For text files
                with open(actual_path, encoding="utf-8", errors="replace") as f:
                    lines = f.readlines()

                if head is not None:
                    lines = lines[: int(head)]
                elif tail is not None:
                    lines = lines[-int(tail) :]

                return "".join(lines)

        elif tool_name == "read_multiple_files":
            paths = params.get("paths", [])
            results = {}
            for file_path in paths:
                try:
                    result = self.handle_filesystem(
                        handler, {"path": file_path}, {"name": "read_file"}
                    )
                    results[file_path] = {"content": result, "error": None}
                except Exception as e:
                    results[file_path] = {"content": None, "error": str(e)}
            return results

        elif tool_name == "list_directory":
            dir_path = params.get("path", "")

            # Normalize path
            if dir_path.startswith("/root/"):
                rel_path = dir_path[6:]
            elif dir_path.startswith("/root"):
                rel_path = dir_path[5:] if len(dir_path) > 5 else ""
            elif dir_path.startswith("/"):
                rel_path = dir_path[1:]
            else:
                rel_path = dir_path

            actual_path = files_root / rel_path if rel_path else files_root

            if not actual_path.exists():
                raise FileNotFoundError(f"Directory not found: {dir_path}")

            entries = []
            for entry in sorted(actual_path.iterdir()):
                prefix = "[DIR]" if entry.is_dir() else "[FILE]"
                entries.append(f"{prefix} {entry.name}")

            return "\n".join(entries)

        elif tool_name == "read_pdf":
            try:
                from pypdf import PdfReader
            except ImportError:
                return [{"error": "pypdf library is not installed"}]

            sources = params.get("sources", [])
            results = []

            pdf_metadata = self.load_api_data("pdf_metadata.json")

            for source in sources:
                path = source.get("path", "")
                # Normalize path
                if not path.startswith("/root"):
                    path = f"/root/{path}"

                # Convert /root path to actual file path
                if path.startswith("/root/"):
                    rel_path = path[6:]
                elif path.startswith("/root"):
                    rel_path = path[5:] if len(path) > 5 else ""
                else:
                    rel_path = path

                actual_path = files_root / rel_path

                if not actual_path.exists():
                    results.append({"path": path, "error": f"PDF not found: {path}"})
                    continue

                try:
                    reader = PdfReader(actual_path)
                    page_count = len(reader.pages)

                    # Extract text from all pages
                    text_parts = []
                    for i, page in enumerate(reader.pages):
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(f"--- Page {i + 1} ---\n{page_text}")

                    full_text = "\n\n".join(text_parts)

                    # Get metadata from PDF or fallback to our metadata file
                    meta = pdf_metadata.get(path, {})
                    pdf_info = reader.metadata
                    if pdf_info:
                        if not meta.get("title") and pdf_info.title:
                            meta["title"] = pdf_info.title
                        if not meta.get("authors") and pdf_info.author:
                            meta["authors"] = [pdf_info.author]

                    results.append(
                        {
                            "path": path,
                            "metadata": meta,
                            "page_count": page_count,
                            "text": full_text,
                        }
                    )
                except Exception as e:
                    results.append({"path": path, "error": f"Error reading PDF: {e}"})

            return results

        else:
            raise ValueError(f"Unknown filesystem tool: {tool_name}")

    def read_docx(self, path: Path) -> str:
        """Extract text from a Word document."""
        try:
            from docx import Document
        except ImportError:
            return "[Error: python-docx library is not installed. Install with: pip install python-docx]"

        doc = Document(path)
        content_parts = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                content_parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.replace("|", "").strip():
                    content_parts.append(row_text)

        return "\n".join(content_parts)

    def read_csv(self, path: Path) -> list[dict]:
        """Read a CSV file and return as list of dicts."""
        with open(path, newline="", encoding="utf-8") as f:
            reader = csv.DictReader(f)
            return list(reader)

    def handle_table_lookup(self, handler: dict, params: dict) -> Any:
        """Handle table lookup from JSON data files."""
        dataset_path = handler.get("dataset", "")
        key_field = handler.get("key_field", "")
        nested_path = handler.get("nested_path", "")

        # Load the dataset
        data = self.load_api_data(dataset_path)

        # Navigate to nested path if specified (e.g., "commodities" or "rates")
        if nested_path and nested_path in data:
            data = data[nested_path]

        # Find the key in params - try common parameter names
        lookup_key = None
        for param_name in [key_field, "nct_id", "paper_id", "dependency", "artifact"]:
            if param_name in params:
                lookup_key = params[param_name]
                break

        if lookup_key is None:
            # Return all data if no key specified (for search operations)
            return data

        # Look up the value
        if lookup_key in data:
            return data[lookup_key]

        # Try with normalized key (remove version suffix for maven)
        base_key = lookup_key.rsplit(":", 1)[0] if ":" in lookup_key else lookup_key
        if base_key in data:
            result = data[base_key].copy()
            result["queried_key"] = lookup_key
            return result

        return {"error": f"Key not found: {lookup_key}"}

    def handle_table_search(self, handler: dict, params: dict, tool: dict) -> Any:
        """Handle search over JSON table data.
        
        Returns matching entries plus some synthetic decoy results.
        Used for trial_searcher, search, list_maven_versions, etc.
        """
        dataset_path = handler.get("dataset", "")
        search_fields = handler.get("search_fields", [])
        result_format = handler.get("result_format", "list")
        include_decoys = handler.get("include_decoys", True)
        max_results = params.get("max_results", params.get("page_size", 10))
        
        # Load the dataset
        data = self.load_api_data(dataset_path)
        
        # Get search query from various possible parameter names
        query = None
        for param_name in ["query", "conditions", "nct_id", "id", "dependency"]:
            if param_name in params:
                query = str(params[param_name])
                break
        
        results = []
        
        # Search for matching entries
        if isinstance(data, dict):
            for key, entry in data.items():
                if query:
                    # Check if query matches key or any searchable field
                    searchable = str(key).lower()
                    if search_fields:
                        for field in search_fields:
                            if field in entry:
                                searchable += " " + str(entry[field]).lower()
                    
                    if query.lower() in searchable:
                        results.append(entry)
                else:
                    results.append(entry)
        elif isinstance(data, list):
            for entry in data:
                if query:
                    searchable = json.dumps(entry).lower()
                    if query.lower() in searchable:
                        results.append(entry)
                else:
                    results.append(entry)
        
        # Add decoy results if configured
        if include_decoys and len(results) > 0:
            decoys = handler.get("decoys", [])
            results.extend(decoys[:max(0, max_results - len(results))])
        
        # Format results
        if result_format == "list":
            return {"results": results[:max_results], "total": len(results)}
        else:
            return results[:max_results]

    def handle_excel_reader(self, handler: dict, params: dict, tool: dict) -> Any:
        """Handle Excel reading operations."""
        tool_name = tool.get("name", "")

        # Get file path
        file_path = params.get("fileAbsolutePath") or params.get("inputPath", "")

        # Map to our data directory
        if file_path.startswith("/root/"):
            rel_path = file_path[6:]
        else:
            rel_path = file_path.lstrip("/")

        actual_path = self.data_path / "files" / "root" / rel_path

        if not actual_path.exists():
            raise FileNotFoundError(f"Excel file not found: {file_path}")

        # Use openpyxl to read the file
        try:
            import openpyxl

            wb = openpyxl.load_workbook(actual_path, data_only=True)
        except ImportError:
            return {"error": "openpyxl not installed"}

        if tool_name in ("excel_describe_sheets", "describe_sheets"):
            sheets = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                sheets.append(
                    {"name": sheet_name, "rows": ws.max_row, "columns": ws.max_column}
                )
            return {"sheets": sheets}

        elif tool_name in ("excel_read_sheet", "excel_read"):
            sheet_name = params.get("sheetName")
            if sheet_name:
                ws = wb[sheet_name]
            else:
                ws = wb.active

            # Read the data
            include_headers = params.get("includeHeaders", True)

            rows = []
            for row in ws.iter_rows(values_only=True):
                rows.append(list(row))

            if include_headers and rows:
                headers = rows[0]
                data = []
                for row in rows[1:]:
                    data.append(dict(zip(headers, row)))
                return {"headers": headers, "data": data, "row_count": len(data)}

            return {"data": rows, "row_count": len(rows)}

        return {"error": f"Unknown Excel tool: {tool_name}"}

    def get_web_corpus_metadata(self) -> dict:
        """Load and cache web corpus metadata."""
        if not self.web_corpus_metadata:
            metadata_path = self.data_path / "web" / "metadata.json"
            if metadata_path.exists():
                with open(metadata_path, encoding="utf-8") as f:
                    data = json.load(f)
                    # Filter out schema/comment fields
                    self.web_corpus_metadata = {
                        k: v for k, v in data.items() if not k.startswith("_")
                    }
        return self.web_corpus_metadata

    def get_web_search_index(self) -> list:
        """Load web search index."""
        index_path = self.data_path / "web" / "search_index.json"
        if index_path.exists():
            with open(index_path, encoding="utf-8") as f:
                data = json.load(f)
                return data.get("entries", [])
        return []

    def get_decoy_urls(self) -> list:
        """Load and cache decoy URLs for search results."""
        if not self.decoy_urls:
            decoy_path = self.data_path / "web" / "decoy_urls.json"
            if decoy_path.exists():
                with open(decoy_path, encoding="utf-8") as f:
                    data = json.load(f)
                    self.decoy_urls = data.get("decoys", [])
        return self.decoy_urls

    def get_decoy_error(self, decoy: dict) -> dict:
        """Generate a realistic error response for a decoy URL."""
        error_type = decoy.get("error_type", "connection_timeout")
        url = decoy.get("url", "")

        error_messages = {
            "connection_timeout": {
                "error": "Connection timed out while trying to reach the server",
                "status": 504,
                "error_code": "ETIMEDOUT",
            },
            "connection_reset": {
                "error": "Connection was reset by the remote server",
                "status": 502,
                "error_code": "ECONNRESET",
            },
            "503_service_unavailable": {
                "error": "Service temporarily unavailable. Please try again later.",
                "status": 503,
                "error_code": "SERVICE_UNAVAILABLE",
            },
            "404_not_found": {
                "error": "The requested page could not be found",
                "status": 404,
                "error_code": "NOT_FOUND",
            },
            "ssl_error": {
                "error": "SSL certificate verification failed",
                "status": 495,
                "error_code": "SSL_ERROR",
            },
        }

        error_info = error_messages.get(
            error_type, error_messages["connection_timeout"]
        )
        return {
            "url": url,
            "success": False,
            **error_info,
        }

    def handle_web_corpus(self, handler: dict, params: dict, tool: dict) -> Any:
        """Handle web corpus operations (Playwright-like fetch)."""
        operation = handler.get("operation", "get_visible_html")
        url = params.get("url", "")

        if not url:
            return {"error": "URL parameter is required", "status": 400}

        # Load corpus metadata
        metadata = self.get_web_corpus_metadata()

        # Check if this is a decoy URL - return realistic error
        decoys = self.get_decoy_urls()
        for decoy in decoys:
            if decoy.get("url") == url:
                return self.get_decoy_error(decoy)

        if operation == "navigate":
            # Check if URL exists in corpus
            if url in metadata:
                entry = metadata[url]
                return {
                    "page_id": entry.get("id", "unknown"),
                    "url": url,
                    "title": entry.get("title", ""),
                    "success": True,
                }
            else:
                return {
                    "error": "URL not found in synthetic web corpus",
                    "status": 404,
                    "url": url,
                }

        elif operation == "get_visible_html":
            if url not in metadata:
                return {
                    "error": "URL not found in synthetic web corpus",
                    "status": 404,
                    "url": url,
                }

            entry = metadata[url]
            html_path = entry.get("html_path", "")
            if not html_path:
                return {
                    "error": "No HTML path configured for URL",
                    "status": 500,
                    "url": url,
                }

            # Load HTML content
            full_html_path = self.data_path / html_path
            if not full_html_path.exists():
                return {"error": "HTML file not found", "status": 500, "url": url}

            with open(full_html_path, encoding="utf-8", errors="replace") as f:
                html_content = f.read()

            return {
                "url": url,
                "html": html_content,
                "title": entry.get("title", ""),
            }

        elif operation == "screenshot":
            # For screenshot, just return metadata about what would be captured
            if url not in metadata:
                return {
                    "error": "URL not found in synthetic web corpus",
                    "status": 404,
                    "url": url,
                }
            entry = metadata[url]
            return {
                "url": url,
                "title": entry.get("title", ""),
                "screenshot": "[Screenshot of page - see HTML content for actual data]",
            }

        else:
            return {
                "error": f"Unknown web_corpus operation: {operation}",
                "status": 400,
            }

    def handle_url_search(self, handler: dict, params: dict, tool: dict) -> Any:
        """Handle URL search operations.

        This performs a simple keyword-based search over the web corpus,
        including both real URLs and decoy URLs. Decoys have a relevance
        penalty applied to make them rank lower than matching real URLs.
        """
        query = params.get("query", "")
        max_results = params.get("max_results", 5)

        if not query:
            return {"error": "Query parameter is required", "results": []}

        # Load search index and decoys
        search_index = self.get_web_search_index()
        decoy_urls = self.get_decoy_urls()

        # Simple keyword matching
        query_lower = query.lower()
        query_terms = query_lower.split()

        def score_entry(entry: dict, is_decoy: bool = False) -> float:
            """Score an entry based on keyword matches."""
            score = 0.0
            searchable_text = " ".join(
                [
                    entry.get("title", ""),
                    entry.get("short_description", ""),
                    " ".join(entry.get("tags", [])),
                    " ".join(entry.get("example_queries", [])),
                ]
            ).lower()

            # Score based on term matches
            for term in query_terms:
                if term in searchable_text:
                    score += 1
                # Bonus for exact phrase match
                if query_lower in searchable_text:
                    score += 2

            # Apply relevance penalty for decoys
            if is_decoy and score > 0:
                penalty = entry.get("relevance_penalty", 0.3)
                score *= 1 - penalty

            return score

        scored_results = []

        # Score real entries
        for entry in search_index:
            score = score_entry(entry, is_decoy=False)
            if score > 0:
                scored_results.append((score, entry, False))

        # Score decoy entries
        for decoy in decoy_urls:
            score = score_entry(decoy, is_decoy=True)
            if score > 0:
                scored_results.append((score, decoy, True))

        # Sort by score descending
        scored_results.sort(key=lambda x: x[0], reverse=True)

        # Build results
        results = []
        for score, entry, is_decoy in scored_results[:max_results]:
            url = entry.get("url", "")
            results.append(
                {
                    "id": entry.get("id", ""),
                    "url": url,
                    "title": entry.get("title", ""),
                    "description": entry.get("short_description", ""),
                }
            )

        if not results:
            return {
                "error": "Search service is temporarily unavailable. Please try again later.",
                "query": query,
            }

        return {
            "query": query,
            "results": results,
            "total_found": len(scored_results),
        }

    def handle_hackernews_story(self, handler: dict, params: dict, tool: dict) -> Any:
        """Handle HackerNews story lookup by ID."""
        story_id = params.get("id", "")

        if not story_id:
            return {"error": "Story ID is required"}

        stories_path = self.data_path / "api" / "hackernews_stories.json"
        if not stories_path.exists():
            return {"error": "HackerNews data not available"}

        with open(stories_path, encoding="utf-8") as f:
            data = json.load(f)

        stories = data.get("stories", {})
        story = stories.get(story_id)

        if not story:
            return {"error": f"Story {story_id} not found"}

        return {
            "id": story["id"],
            "title": story["title"],
            "url": story["url"],
            "author": story["author"],
            "points": story["points"],
            "num_comments": story["num_comments"],
            "age": story["age"],
            "discussion_url": story["discussion_url"],
        }

    def handle_wikipedia_search(self, handler: dict, params: dict, tool: dict) -> Any:
        """Handle Wikipedia article search."""
        query = params.get("query", "")

        if not query:
            return {"error": "Search query is required", "results": []}

        articles_path = self.data_path / "api" / "wikipedia_articles.json"
        if not articles_path.exists():
            return {"error": "Wikipedia data not available", "results": []}

        with open(articles_path, encoding="utf-8") as f:
            data = json.load(f)

        query_lower = query.lower()
        query_terms = query_lower.split()

        results = []
        for article in data.get("articles", []):
            searchable = " ".join(
                [article.get("title", ""), article.get("summary", "")]
                + article.get("keywords", [])
            ).lower()

            if any(term in searchable for term in query_terms):
                results.append(
                    {
                        "title": article["title"],
                        "url": article["url"],
                        "summary": article["summary"],
                    }
                )

        if not results:
            return {
                "error": "Wikipedia search is temporarily unavailable. Please try again later.",
                "query": query,
            }

        return {"query": query, "results": results}

    def load_api_data(self, dataset_path: str) -> dict:
        """Load a JSON dataset from the API data directory."""
        # Handle different path formats
        if dataset_path.startswith("data/"):
            rel_path = dataset_path[5:]  # Remove 'data/' prefix
        else:
            rel_path = dataset_path

        # Try in api/ subdirectory first
        full_path = self.data_path / "api" / Path(rel_path).name
        if not full_path.exists():
            full_path = self.data_path / rel_path

        if not full_path.exists():
            return {}

        with open(full_path, encoding="utf-8") as f:
            return json.load(f)


def create_server(
    config_path: Path = DEFAULT_CONFIG_PATH,
    data_path: Path = DEFAULT_DATA_PATH,
    logs_path: Path = DEFAULT_LOGS_PATH,
    port: int = 8765,
    log_stub_calls: bool = True,
) -> HTTPServer:
    """Create and configure the HTTP server.
    
    Args:
        config_path: Path to servers.json configuration
        data_path: Path to synthetic data directory
        logs_path: Path to logs directory
        port: Port to listen on
        log_stub_calls: If True, log calls to stub handlers for review
    """
    # Load server configuration
    with open(config_path, encoding="utf-8") as f:
        servers_config = json.load(f)

    # Configure the handler class
    SyntheticMCPHandler.servers_config = servers_config
    SyntheticMCPHandler.data_path = data_path
    
    # Set up stub logging if enabled
    if log_stub_calls:
        stub_log_path = logs_path / "stub_calls.json"
        SyntheticMCPHandler.stub_logger = StubCallLogger(stub_log_path)
    else:
        SyntheticMCPHandler.stub_logger = None

    server = HTTPServer(("", port), SyntheticMCPHandler)
    return server


def main():
    """Run the synthetic MCP server."""
    import argparse

    parser = argparse.ArgumentParser(description="Synthetic HTTP MCP Server")
    parser.add_argument("--port", type=int, default=8765, help="Port to listen on")
    parser.add_argument(
        "--config", type=Path, default=DEFAULT_CONFIG_PATH, help="Path to servers.json"
    )
    parser.add_argument(
        "--data", type=Path, default=DEFAULT_DATA_PATH, help="Path to data directory"
    )
    parser.add_argument(
        "--logs", type=Path, default=DEFAULT_LOGS_PATH, help="Path to logs directory"
    )
    parser.add_argument(
        "--no-stub-logging",
        action="store_true",
        help="Disable logging of stub handler calls",
    )
    args = parser.parse_args()

    server = create_server(
        config_path=args.config,
        data_path=args.data,
        logs_path=args.logs,
        port=args.port,
        log_stub_calls=not args.no_stub_logging,
    )

    print(f"Synthetic MCP Server starting on port {args.port}")
    print(f"Config: {args.config}")
    print(f"Data: {args.data}")
    print(f"Servers: {list(SyntheticMCPHandler.servers_config.keys())}")
    if not args.no_stub_logging:
        print(f"Stub call logging: {args.logs / 'stub_calls.json'}")

    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print("\nShutting down...")
        server.shutdown()


if __name__ == "__main__":
    main()
